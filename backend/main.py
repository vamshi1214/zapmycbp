from fastapi import FastAPI, HTTPException, UploadFile, File
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse
from pydantic import BaseModel
from typing import List, Optional, Union
import google.generativeai as genai
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn, nsdecls
from docx.oxml import ns
import datetime
import shutil
from dotenv import load_dotenv
from fastapi.staticfiles import StaticFiles
import pymongo
from pymongo.server_api import ServerApi
from collections import OrderedDict
import uuid
import imghdr
from PIL import Image
import asyncio
import logging
import json

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

load_dotenv()

app = FastAPI()

# Configure CORS
app.add_middleware(
    CORSMiddleware,
    allow_origins=["http://localhost:8080"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Get the absolute path to the current directory
current_dir = os.path.dirname(os.path.abspath(__file__))

# Create directories if they don't exist
reports_dir = os.path.join(current_dir, "reports")
uploads_dir = os.path.join(current_dir, "uploads")
os.makedirs(reports_dir, exist_ok=True)
os.makedirs(uploads_dir, exist_ok=True)

# Constants for file validation
MAX_FILE_SIZE = 5 * 1024 * 1024  # 5MB
ALLOWED_EXTENSIONS = {'png', 'jpg', 'jpeg', 'gif'}
MAX_IMAGES_PER_USER = 10

# Session tracking
active_sessions = {}

# Add this after your app initialization
app.mount("/uploads", StaticFiles(directory="uploads"), name="uploads")

def cleanup_session_images(session_id: str):
    """Clean up images for a specific session"""
    try:
        if session_id in active_sessions:
            for filename in active_sessions[session_id]:
                file_path = os.path.join(uploads_dir, filename)
                if os.path.exists(file_path):
                    os.remove(file_path)
                    logger.info(f"Cleaned up session image: {filename}")
            del active_sessions[session_id]
    except Exception as e:
        logger.error(f"Error cleaning up session images: {str(e)}")

@app.post("/api/start-session")
async def start_session():
    """Start a new session and return session ID"""
    session_id = str(uuid.uuid4())
    active_sessions[session_id] = []
    return {"sessionId": session_id}

@app.post("/api/end-session/{session_id}")
async def end_session(session_id: str):
    """End a session and cleanup its images"""
    cleanup_session_images(session_id)
    return {"status": "success"}

def generate_unique_filename(original_filename: str) -> str:
    """Generate a unique filename using UUID"""
    ext = os.path.splitext(original_filename)[1]
    return f"{uuid.uuid4()}{ext}"

def validate_image(file: UploadFile):
    """Validate image file size and type"""
    if file.content_type not in [f'image/{ext}' for ext in ALLOWED_EXTENSIONS]:
        raise HTTPException(status_code=400, detail="Invalid file type. Only images are allowed.")
    
    file.file.seek(0, 2)  # Seek to end of file
    size = file.file.tell()  # Get file size
    file.file.seek(0)  # Reset file pointer
    
    if size > MAX_FILE_SIZE:
        raise HTTPException(status_code=400, detail="File too large. Maximum size is 5MB.")

@app.post("/api/upload-image")
async def upload_image(file: UploadFile = File(...), session_id: str = None):
    try:
        # Validate image
        validate_image(file)

        # Generate unique filename
        unique_filename = generate_unique_filename(file.filename)
        file_path = os.path.join(uploads_dir, unique_filename)

        # Save the uploaded file
        with open(file_path, "wb") as buffer:
            shutil.copyfileobj(file.file, buffer)

        # Verify the saved image is valid
        try:
            with Image.open(file_path) as img:
                img.verify()  # Verify it's a valid image
        except Exception as e:
            os.remove(file_path)  # Remove invalid image
            raise HTTPException(status_code=400, detail="Invalid or corrupted image file")

        # Track the image in the session
        if session_id and session_id in active_sessions:
            active_sessions[session_id].append(unique_filename)

        return {"filename": unique_filename}
    except Exception as e:
        logger.error(f"Error uploading image: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))

class TeamMember(BaseModel):
    name: str
    rollNumber: str
    gender: str

class ProjectResult(BaseModel):
    resultImages: Optional[List[str]] = None
    codeOutput: Optional[str] = None
    aiGeneratedContent: Optional[Union[str, bool]] = None

class ProjectData(BaseModel):
    projectDescription: str
    projectCode: str
    department: str
    mainProfessor: str
    mainProfessor_designation: str
    professorDepartment: str
    secondaryProfessor: Optional[str] = None
    secondaryProfessor_designation: Optional[str] = None
    course: str
    teamMembers: List[TeamMember]
    result: Optional[ProjectResult] = None

# Function to fetch professor details from MongoDB
'''def get_professor_details_from_db(professor_name: str):
    professor_document = professor_collection.find_one({"name": professor_name})
    if professor_document:
        return professor_document.get("designation")
    else:
        return None '''

def add_page_border(section):
    """
    Add a border around the page by modifying the section's XML.
    """
    sectPr = section._sectPr
    if sectPr is None:
        sectPr = OxmlElement('w:sectPr')
        section._element.append(sectPr)

    pgBorders = OxmlElement("w:pgBorders")
    pgBorders.set(ns.qn('w:offsetFrom'), 'page')  # Ensure border is relative to page edge

    # Set border attributes
    for border in ["top", "left", "bottom", "right"]:
        border_element = OxmlElement(f"w:{border}")
        border_element.set(ns.qn("w:val"), "single")  # Single line border
        border_element.set(ns.qn("w:sz"), "6")  # Border size (in eighths of a point)
        border_element.set(ns.qn("w:space"), "24")  # Space between border and content in points
        border_element.set(ns.qn("w:color"), "000000")  # Black color
        pgBorders.append(border_element)

    sectPr.append(pgBorders)

def add_page_number(paragraph):
    """
    Add a centered page number to a paragraph.
    """
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    page_num_run = paragraph.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'PAGE'
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    page_num_run._r.append(fldChar1)
    page_num_run._r.append(instrText)
    page_num_run._r.append(fldChar2)

def format_text_content(doc, text_content):
    """
    Format text content with proper styling for bullets, bold text, and headers.
    
    Args:
        doc: The Document object
        text_content: The text content to format
    """
    if not text_content:
        return
    

        
    for line in text_content.splitlines():
        # Skip empty lines
        if not line.strip():
            continue

        # Bold section titles (if they start and end with **)
        if line.startswith("**") and line.endswith("**"):
            paragraph = doc.add_paragraph()
            bold_text = paragraph.add_run(line.strip("**"))
            bold_text.bold = True
            
        # Handle section titles with asterisks (e.g., *Video Processing:*)
        elif line.startswith("*") and line.endswith("*") and ":" in line:
            paragraph = doc.add_paragraph()
            title_text = line.strip("*")
            run = paragraph.add_run(title_text)
            run.bold = True
            run.font.size = Pt(12)

        # Handle top-level bullet points (e.g., `•`)
        elif line.startswith("•"):
            clean_line = line.lstrip("• ").strip()

            # Check for bold text within bullet points
            if "**" in clean_line:
                paragraph = doc.add_paragraph(style="List Bullet")
                parts = clean_line.split("**")
                for i, part in enumerate(parts):
                    run = paragraph.add_run(part)
                    if i % 2 == 1:  # Apply bold to parts between `**`
                        run.bold = True
            else:
                doc.add_paragraph(clean_line, style="List Bullet")

        # Handle sub-bullet points with asterisks at beginning (e.g., `* Extract frames`)
        elif line.strip().startswith("* "):
            clean_line = line.strip().lstrip("* ").strip()

            # Check for bold text within sub-bullets
            if "**" in clean_line:
                paragraph = doc.add_paragraph(style="List Bullet 2")
                parts = clean_line.split("**")
                for i, part in enumerate(parts):
                    run = paragraph.add_run(part)
                    if i % 2 == 1:  # Apply bold to parts between `**`
                        run.bold = True
            else:
                doc.add_paragraph(clean_line, style="List Bullet 2")

        # Handle numbered lists (e.g., `1.`)
        elif line[0].isdigit() and len(line) > 1 and line[1] == ".":
            parts = line.split(". ", 1)  # Split at ". "
            if len(parts) > 1:
                number = parts[0]
                rest_of_line = parts[1]

                paragraph = doc.add_paragraph()
                paragraph.add_run(number + ". ")  # Number is NOT bold

                # Handle bolding within the description
                if "**" in rest_of_line:
                    bold_parts = rest_of_line.split("**")
                    for i, part in enumerate(bold_parts):
                        run = paragraph.add_run(part)
                        if i % 2 == 1:  # Odd indices are the bold parts
                            run.bold = True
                else:
                    paragraph.add_run(rest_of_line)

        # Handle regular top-level bullet points
        elif line.startswith("- "):
            clean_line = line.lstrip("- ").strip()

            # Check for bold text within the bullet point
            if "**" in clean_line:
                paragraph = doc.add_paragraph(style="List Bullet")
                parts = clean_line.split("**")
                for i, part in enumerate(parts):
                    run = paragraph.add_run(part)
                    if i % 2 == 1:  # Apply bold to parts between `**`
                        run.bold = True
            else:
                doc.add_paragraph(clean_line, style="List Bullet")

        # Handle sub-bullet points (indented bullets)
        elif line.startswith("    - ") or line.startswith("\t- "):  # Four spaces or tab for indentation
            clean_line = line.lstrip("\t ").lstrip("- ").strip()
            paragraph = doc.add_paragraph(clean_line, style="List Bullet 2")
        elif line.startswith("  - "):  # Sub-bullets
            clean_line = line.lstrip("  - ").strip()

            # Check for bold text within the sub-bullet
            if "**" in clean_line:
                paragraph = doc.add_paragraph(style="List Bullet 2")
                parts = clean_line.split("**")
                for i, part in enumerate(parts):
                    run = paragraph.add_run(part)
                    if i % 2 == 1:  # Apply bold to parts between `**`
                        run.bold = True
            else:
                doc.add_paragraph(clean_line, style="List Bullet 2")

        elif line.startswith("        -"):  # Sub-bullets
            clean_line = line.lstrip("        -").strip()

            # Check for bold text within the sub-bullet
            if "**" in clean_line:
                paragraph = doc.add_paragraph(style="List Bullet 2")
                parts = clean_line.split("**")
                for i, part in enumerate(parts):
                    run = paragraph.add_run(part)
                    if i % 2 == 1:  # Apply bold to parts between `**`
                        run.bold = True
            else:
                doc.add_paragraph(clean_line, style="List Bullet 2")

        # Handle section titles (e.g., `###` for level 3 headers)
        elif line.startswith("###"):
            # Create a new paragraph for the header
            paragraph = doc.add_paragraph()
            header_text = line.lstrip("###").strip()
            run = paragraph.add_run(header_text)
            run.bold = True  # You can adjust this to apply a different style (e.g., font size)
            paragraph.style = "Heading 3"
            
        # Regular plain text (non-bulleted)
        else:
            paragraph = doc.add_paragraph(line)
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

@app.post("/api/generate-report")
async def generate_report(data: ProjectData, session_id: str = None):
    try:
        # Add default values for required fields if they're empty
        if not data.department or data.department == "":
            data.department = "Computer Science"
            
        if not data.professorDepartment or data.professorDepartment == "":
            data.professorDepartment = "Computer Science & Engineering"
            
        if not data.course or data.course == "":
            data.course = "Technical Course"
        
        # Check if GEMINI_API_KEY is set
        api_key = os.getenv("GEMINI_API_KEY")
        if not api_key:
            raise HTTPException(
                status_code=500,
                detail="GEMINI_API_KEY environment variable is not set"
            )

        # Configure Gemini API
        genai.configure(api_key=api_key)
        model = genai.GenerativeModel('gemini-1.5-flash')
        chat = model.start_chat(history=[])
        
        # Generate AI content if needed
        if data.result and isinstance(data.result.aiGeneratedContent, bool) and data.result.aiGeneratedContent:
            try:
                # Generate AI analysis of the code and output
                prompt = f"""Analyze the following code and its output, providing insights about:
                1. The code's functionality and performance
                2. Key patterns or interesting aspects in the output
                3. Potential improvements or optimizations
                4. Any notable technical achievements

                Code: {data.projectCode}
                Output: {data.result.codeOutput if data.result.codeOutput else 'No output provided'}
                """
                
                ai_response = chat.send_message(prompt)
                # Replace the boolean with the generated text
                data.result.aiGeneratedContent = ai_response.text if ai_response.text else "AI analysis could not be generated."
            except Exception as e:
                print(f"Error generating AI content: {e}")
                data.result.aiGeneratedContent = "Error generating code and output analysis."
        
        # Get the absolute path to the current directory
        current_dir = os.path.dirname(os.path.abspath(__file__))

        # Create reports directory if it doesn't exist
        reports_dir = os.path.join(current_dir, "reports")
        os.makedirs(reports_dir, exist_ok=True)

        # Define paths
        report_path = os.path.join(reports_dir, "report.docx")
        logo_path = os.path.join(current_dir, "logo.jpg")

        # Verify logo exists
        if not os.path.exists(logo_path):
            raise HTTPException(
                status_code=500,
                detail="Logo file not found"
            )

        # Your existing Python script logic here
        #### Date
        # Get the current year and format it as "YYYY-YYYY+1"
        current_year = datetime.datetime.now().year
        formatted_year = f"{current_year}-{current_year + 1}"

        # Extract data from the request payload
        team_members = data.teamMembers
        project_code = data.projectCode
        project_description = data.projectDescription
        department = data.department
        main_professor = data.mainProfessor
        main_professor_designation = data.mainProfessor_designation
        professor_department = data.professorDepartment
        secondary_professor = data.secondaryProfessor
        secondary_professor_designation = data.secondaryProfessor_designation
        course = data.course

        # Combine into lists for easier processing
        n_s = [member.name for member in team_members]
        r_s = [member.rollNumber for member in team_members]
        g_s = [member.gender for member in team_members]

        chat = model.start_chat(history=[])

        def get_text(response):
            return response.text if response.text else ""

        ##for title
        title_text = get_text(chat.send_message(
            f"Using the provided description and code, get me a title for the code with Uppercase letters. Description: {project_description}. Code: {project_code}"))

        try:
            abstract = get_text(chat.send_message(f"Using the provided description and code, write a concise 400-word abstract summarizing the project. Do not include any titles or headings in your response. Replace any generic terms like 'the project'  with the specific project title wherever applicable. Don't include any conclusion. Description: {project_description}. Code: {project_code}"))
            introduction = get_text(chat.send_message(f""""Compose a compelling and informative project overview of approximately 350 words, designed to immediately engage the reader and provide a comprehensive understanding of the project's scope and significance. This overview should be structured as follows:

            1. **Context and Motivation:** Begin by establishing the context of the research or problem being addressed. Clearly explain the motivation behind the project and why it is important.
            2. **Objectives and Goals:** Explicitly state the project's objectives and the intended outcomes. What specific goals are you trying to achieve?
            3. **Methodology and Approach:** Describe the key methodologies, techniques, or technologies that will be used to achieve the project's objectives. Provide a high-level overview of the project's workflow or stages.
            4. **Potential Impact and Contributions:** Discuss the potential impact of the project on the relevant field or area of study. What new knowledge, solutions, or insights are expected to emerge?
            5. **Concluding Statement:** Briefly summarize the project's overall significance and reiterate its potential contributions.

            Crucially, this overview must *not* include the word "Introduction" as a heading, or any other section headings. It should flow seamlessly as a single, cohesive piece of text.

            Base your response on the following description and code:

            Description: {project_description}
            Code: {project_code}"""))

            conclusion = get_text(chat.send_message(f"""Compose a compelling and informative concluding summary of approximately 300 words, designed to provide a strong sense of closure and highlight the project's overall impact. This summary should be structured as follows:

            1. **Summary of Outcomes:** Briefly recap the project's main objectives and summarize the key results or outcomes achieved.
            2. **Significance of Findings:** Discuss the significance of these findings in the context of the research area or problem being addressed. What new insights or knowledge have been gained?
            3. **Implications and Impact:** Explore the broader implications of the project's outcomes. What are the potential applications, real-world impacts, or future research directions that stem from this work?
            4. **Limitations and Future Work (Optional):** Briefly acknowledge any limitations of the project and suggest potential avenues for future research or improvement.
            5. **Concluding Remarks:** Offer a concise concluding statement that reinforces the project's overall contribution and significance.

            Crucially, this summary must *not* include the word "Conclusion" as a heading, or any other section headings. It should flow seamlessly as a single, cohesive piece of text.

            Base your response on the following description and code:

            Description: {project_description}
            Code: {project_code}"""))

            objectives = get_text(chat.send_message(f"Analyze the following code and provide a concise, bullet-point summary of its objectives, targeting approximately 20 bullet points. Prioritize clarity, conciseness, and the use of short, direct sentences. Group related objectives under short, descriptive *side headings* if it improves readability and organization. Absolutely *do not* use a main heading for the entire summary. Description: {project_description}, Code: {project_code}"))
        
            methodology = get_text(chat.send_message(f"""Provide a detailed explanation of the methodology employed in this project, based on the provided description and code. This explanation should cover:

            *   The specific steps or stages involved in the project.
            *   The processes and procedures used at each stage.
            *   The techniques, algorithms, or tools applied.
            *   The overall approach or strategy adopted.

            Use side headings to organize the explanation into logical sections, but do *not* use a main heading for the entire response.

            Description: {project_description}
            Code: {project_code}"""))
            
        except Exception as e:
            print(f"Error generating content: {e}")
            abstract = introduction = conclusion = objectives = methodology = ""

        def make_table_invisible(table):
            """Helper function to make table borders invisible"""
            tbl = table._element
            tblPr = tbl.xpath('w:tblPr')
            if not tblPr:
                tblPr = OxmlElement('w:tblPr')
                tbl.insert(0, tblPr)
            else:
                tblPr = tblPr[0]

            # Create border element
            tblBorders = OxmlElement('w:tblBorders')

            # Add all border types
            for border_type in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
                border = OxmlElement(f'w:{border_type}')
                border.set(qn('w:val'), 'none')
                border.set(qn('w:sz'), '0')
                border.set(qn('w:space'), '0')
                border.set(qn('w:color'), 'auto')
                tblBorders.append(border)

            # Remove any existing borders
            existing_borders = tblPr.find(qn('w:tblBorders'))
            if existing_borders is not None:
                tblPr.remove(existing_borders)

            # Add new border settings
            tblPr.append(tblBorders)

        def _add_results_section(doc, result, title_style, project_code):
            """Helper function to add the Results section to the document"""
            section = doc.sections[-1]
            add_page_border(section)

            # Add Results title
            paragraph = doc.add_paragraph()
            paragraph.style = title_style
            paragraph.add_run("RESULTS").bold = True

            # Add result images if they exist
            if result and result.resultImages and len(result.resultImages) > 0:
                # Create a table for images with 2 columns
                image_table = doc.add_table(rows=0, cols=2)
                image_table.autofit = False
                image_table.columns[0].width = Inches(3)
                image_table.columns[1].width = Inches(3)
                
                # Add images in pairs
                for i in range(0, len(result.resultImages), 2):
                    row = image_table.add_row()
                    # First image
                    try:
                        img_path = os.path.join(uploads_dir, result.resultImages[i])
                        cell = row.cells[0]
                        paragraph = cell.paragraphs[0]
                        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        run = paragraph.add_run()
                        run.add_picture(img_path, width=Inches(2.5))
                    except Exception as e:
                        print(f"Error adding image {result.resultImages[i]}: {str(e)}")
                    
                    # Second image (if exists)
                    if i + 1 < len(result.resultImages):
                        try:
                            img_path = os.path.join(uploads_dir, result.resultImages[i + 1])
                            cell = row.cells[1]
                            paragraph = cell.paragraphs[0]
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                            run = paragraph.add_run()
                            run.add_picture(img_path, width=Inches(2.5))
                        except Exception as e:
                            print(f"Error adding image {result.resultImages[i + 1]}: {str(e)}")
                
                # Add space after images
                doc.add_paragraph()

            # Add code output if it exists
            if result and result.codeOutput:
                # Add a heading for code output
                output_heading = doc.add_paragraph()
                output_heading.add_run("Code Output:").bold = True
                
                # Add the code output
                code_output = doc.add_paragraph()
                code_output.add_run(result.codeOutput)

            # Add AI content if it exists
            if result and result.aiGeneratedContent:
                # Add a heading for AI analysis
                ai_heading = doc.add_paragraph()
                ai_heading.add_run("Code & Output Analysis:").bold = True
                
                # Format the AI content using our new function
                format_text_content(doc, str(result.aiGeneratedContent))

            # Add page break after Results section
            doc.add_page_break()

        def create_project_report(
            abstract,
            introduction,
            objectives,
            methodology,
            project_code,
            result,
            conclusion,
            department,
        ):
            from docx.shared import Inches

            # Create a new document
            doc = Document()

            # Add page borders to the first section
            add_page_border(doc.sections[0])

            # Add page numbers to the footer of all sections
            for section in doc.sections:
                footer = section.footer
                paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
                add_page_number(paragraph)

            # Set default font to Times New Roman
            style = doc.styles["Normal"]
            font = style.font
            font.name = "Times New Roman"
            font.size = Pt(12)

            # Title section
            title = doc.add_paragraph()
            title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            run = title.add_run("A Course Based Project Report on\n")
            run.font.size = Pt(14)
            run.font.color.rgb = RGBColor(0, 0, 139)

            run = title.add_run(title_text)
            run.bold = True
            run.font.size = Pt(18)
            run.font.color.rgb = RGBColor(255, 0, 0)

            run = title.add_run("Submitted to the\n")
            run.font.size = Pt(14)
            run.font.color.rgb = RGBColor(0, 0, 139)

            run = title.add_run("Department of " + professor_department + "\n")
            run.bold = True
            run.font.size = Pt(16)
            run.font.color.rgb = RGBColor(0, 0, 139)

            run = title.add_run(
                "in partial fulfilment of the requirements for the completion of course\n"
            )
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor(0, 0, 139)

            run = title.add_run(course + "\n\n")
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor(0, 0, 0)

            run = title.add_run("BACHELOR OF TECHNOLOGY\n")
            run.bold = True
            run.font.size = Pt(14)
            run.font.color.rgb = RGBColor(255, 0, 0)

            run = title.add_run("in\n")
            run.bold = True
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor(0, 0, 0)

            run = title.add_run("Department of " + professor_department)
            run.bold = True
            run.font.size = Pt(16)
            run.font.color.rgb = RGBColor(255, 0, 0)

            # Submitted by section (Combined alignments)
            submitted_by = doc.add_paragraph()
            submitted_by.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # Center "Submitted by"
            run = submitted_by.add_run("Submitted by\n")
            run.font.size = Pt(14)
            run.font.color.rgb = RGBColor(0, 0, 0)

            students = []
            for member in team_members:
                students.append([member.name, member.rollNumber])

            table = doc.add_table(rows=len(students), cols=2)
            table.autofit = False
            table.columns[0].width = Pt(250)
            table.columns[1].width = Pt(100)
            make_table_invisible(table)

            # Indent the table to the right (adjust Pt value as needed)
            # 1 inch indentation

            for i, (name, roll_no) in enumerate(students):
                row_cells = table.rows[i].cells
                # Add spaces before the name:
                padded_name = "                " + name  # Six spaces
                padded_roll_no = "                   " + roll_no
                row_cells[0].text = padded_name
                row_cells[1].text = padded_roll_no

                for cell in row_cells:
                    paragraph = cell.paragraphs[0]
                    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT  # Right-align cell content
                    for run in paragraph.runs:
                        run.font.size = Pt(12)
                        run.font.color.rgb = RGBColor(0, 0, 139)

            # Under the guidance of section
            guidance = doc.add_paragraph()
            guidance.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            run = guidance.add_run("Under the guidance of\n")
            run.font.size = Pt(14)
            run.font.color.rgb = RGBColor(0, 0, 0)

            run = guidance.add_run(main_professor + "\n")
            run.bold = True
            run.font.size = Pt(14)
            run.font.color.rgb = RGBColor(0, 0, 139)

            run = guidance.add_run(
                main_professor_designation
                + ", Department of "
                + professor_department
                + " VNRVJIET"
            )
            run.bold = True
            run.font.size = Pt(14)
            run.font.color.rgb = RGBColor(0, 0, 0)

            image_placeholder = doc.add_paragraph()
            image_placeholder.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            # Add the image with the correct method
            image_placeholder.add_run().add_picture(
                "logo.jpg", width=Inches(1), height=Inches(1)
            )

            header = doc.add_paragraph()
            header.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            run = header.add_run(
                "VALLURUPALLI NAGESWARA RAO VIGNANA JYOTHI INSTITUTE OF ENGINEERING AND TECHNOLOGY\n"
            )
            run.bold = True
            run.font.size = Pt(14)
            run.font.color.rgb = RGBColor(0, 0, 139)

            run = header.add_run(
                "An Autonomous Institute, NAAC Accredited with 'A++' Grade, NBA Accredited for CE, EEE, ME, ECE, CSE, EIE, IT B. Tech Courses, Approved by AICTE, New Delhi, Affiliated to JNTUH, Recognized as 'College with Potential for Excellence' by UGC, ISO 9001:2015 Certified, QS I GUAGE Diamond Rated\n"
            )
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(255, 0, 0)

            run = header.add_run(
                "Vignana Jyothi Nagar, Pragathi Nagar, Nizampet(SO), Hyderabad-500090, TS, India\n"
            )
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0, 0, 139)

            run = header.add_run("Department of " + professor_department)
            run.bold = True
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor(255, 0, 0)

            #####################################Add a page break########################################################################
            doc.add_page_break()
            # Certificate page
            section = doc.sections[-1]
            add_page_border(section)

            header = doc.add_paragraph()
            header.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            run = header.add_run(
                "VALLURUPALLI NAGESWARA RAO VIGNANA JYOTHI INSTITUTE OF ENGINEERING AND TECHNOLOGY\n"
            )
            run.bold = True
            run.font.size = Pt(14)
            run.font.color.rgb = RGBColor(0, 0, 139)

            run = header.add_run(
                "An Autonomous Institute, NAAC Accredited with 'A++' Grade, NBA Accredited for CE, EEE, ME, ECE, CSE, EIE, IT B. Tech Courses, Approved by AICTE, New Delhi, Affiliated to JNTUH, Recognized as 'College with Potential for Excellence' by UGC, ISO 9001:2015 Certified, QS I GUAGE Diamond Rated\n"
            )
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(255, 0, 0)

            run = header.add_run(
                "Vignana Jyothi Nagar, Pragathi Nagar, Nizampet(SO), Hyderabad-500090, TS, India\n\n"
            )
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0, 0, 139)

            run = header.add_run("Department of " + professor_department + "\n\n")
            run.bold = True
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor(255, 0, 0)

            # Leave space for image
            from docx.shared import Inches

            # Add a new paragraph for the image
            image_placeholder = doc.add_paragraph()
            image_placeholder.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            # Add the image with the correct method
            image_placeholder.add_run().add_picture(
                "logo.jpg", width=Inches(1), height=Inches(1)
            )  # Adjust width as necessary

            # Optionally, you can set additional formatting to the image's caption if needed

            run = header.add_run("CERTIFICATE\n\n")
            run.bold = True
            run.font.size = Pt(15)
            run.font.color.rgb = RGBColor(0, 128, 0)
            run.underline = True

            certificate = doc.add_paragraph()
            certificate.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            certificate.space_after = Pt(12)  # Add space after the paragraph

            run = certificate.add_run("This is to certify that the project report entitled ")
            run.font.size = Pt(12)

            run = certificate.add_run('"' + title_text + '"')
            run.font.color.rgb = RGBColor(0, 0, 139)
            run.bold = True
            run.font.size = Pt(12)

            run = certificate.add_run(
                "is a bonafide work done under our supervision and is being submitted by "
            )
            run.font.size = Pt(12)

            students = [
                f"{'Mr.' if gender == 'm' else 'Miss.'} {name} ({roll})"
                for name, roll, gender in zip(n_s, r_s, g_s)
            ]

            for i, student in enumerate(students):
                run = certificate.add_run(student)
                run.font.size = Pt(12)
                run.bold = True
                run.font.color.rgb = RGBColor(0, 128, 0)
                if i < len(students) - 1:
                    run = certificate.add_run(", ")
                    run.bold = True
                    run.font.size = Pt(12)

            run = certificate.add_run(
                "in partial fulfillment for the award of the degree of "
            )
            run.font.size = Pt(12)

            run = certificate.add_run("Bachelor of Technology ")
            run.font.size = Pt(12)
            run.bold = True
            run.font.color.rgb = RGBColor(0, 0, 255)  # Blue color

            run = certificate.add_run(
                "in" + department + ", "
            )
            run.font.size = Pt(12)

            run = certificate.add_run(
                "of the VNR VJIET, Hyderabad during the academic year "
                + formatted_year
                + ".\n\n\n\n"
            )
            run.font.size = Pt(12)

            # After acknowledgment, before TOC
            # Determine HOD based on department
            if department == "Computer Science & Engineering":
                hod_name = "Dr. V. Baby"
            elif department == "Electrical and Electronics Engineering":
                hod_name = "Dr.V. Ramesh Babu"
            elif department == "Electronics and Communication Engineering":
                hod_name = "Dr L Padma Sree"
            elif department == "Mechanical Engineering":
                hod_name = "Dr. B.V.R. Ravi Kumar"
            elif department == "Electronics and Instrumentation Engineering":
                hod_name = "Dr. S. Pranavanand"
            elif department == "Civil Engineering":
                hod_name = "Dr. K. Ramujee"
            elif department == "Automobile Engineering":
                hod_name = "Dr.Shaik Amjad"
            elif department == "Artificial Intelligence & Data Science":
                hod_name = "Dr.T.Sunil Kumar"
            elif department == "CSE-Cyber Security":
                hod_name = "Dr.T.Sunil Kumar"
            elif department == "CSE-Data Science":
                hod_name = "Dr.T.Sunil Kumar"
            elif department == "Computer Science and Business Systems":
                hod_name = "Dr. V. Baby"
            elif department == "CSE-AIML":
                hod_name = "Dr.Sagar Yeruva"
            elif department == "CSE-IoT":
                hod_name = "Dr.Sagar Yeruva"
            elif department == "Information Technology":
                hod_name = "Dr N Mangathayaru"
            else:
                hod_name = "Department Head"  # Default value

            # Create signatures table
            signatures = doc.add_table(rows=4, cols=2)  # 4 rows now
            signatures.autofit = False
            signatures.style.paragraph_format.space_before = Pt(12)

            # Set column widths
            signatures.columns[0].width = Inches(2.5)
            signatures.columns[1].width = Inches(2.5)

            # --- Project Guide ---
            cell = signatures.cell(0, 0)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            run = p.add_run(main_professor)
            run.bold = True
            run.font.color.rgb = RGBColor(255, 0, 0)
            run.font.size = Pt(12)

            cell = signatures.cell(1, 0)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            run = p.add_run(main_professor_designation)
            run.font.size = Pt(11)

            cell = signatures.cell(2, 0)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            run = p.add_run("Dept of " + professor_department)
            run.font.color.rgb = RGBColor(255, 0, 0)
            run.font.size = Pt(11)

            # --- HOD ---
            cell = signatures.cell(0, 1)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            run = p.add_run("          " + hod_name)
            run.font.color.rgb = RGBColor(255, 0, 0)
            run.bold = True
            run.font.size = Pt(12)

            cell = signatures.cell(1, 1)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            run = p.add_run("            " + "Professor & HOD")
            run.font.size = Pt(11)

            cell = signatures.cell(2, 1)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            run = p.add_run("            " + "Dept of " + department)
            run.font.color.rgb = RGBColor(255, 0, 0)
            run.font.size = Pt(11)

            ########################3rd page#######################################################################################################
            doc.add_page_break()
            section = doc.sections[-1]
            add_page_border(section)

            # Title Section
            title = doc.add_paragraph()
            title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            run = title.add_run("Course based Projects Reviewer\n")
            run.font.size = Pt(14)
            run.bold = True
            run.font.color.rgb = RGBColor(0, 0, 139)

            run = title.add_run(
                "VALLURUPALLI NAGESWARA RAO VIGNANA JYOTHI INSTITUTE OF ENGINEERING AND TECHNOLOGY\n"
            )
            run.font.size = Pt(12)
            run.bold = True
            run.font.color.rgb = RGBColor(255, 0, 0)

            run = title.add_run(
                "An Autonomous Institute, NAAC Accredited with 'A++' Grade,\nVignana Jyothi Nagar, Pragathi Nagar, Nizampet(SO), Hyderabad-500090, TS, India\n\n"
            )
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0, 0, 139)

            run = title.add_run("Department of " + professor_department + "\n")
            run.bold = True
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor(255, 0, 0)

            # Image Placeholder
            doc.add_picture(
                "logo.jpg", width=Inches(1), height=Inches(1)
            )  # Adjust path and width
            doc.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            # Declaration Section
            declaration_title = doc.add_paragraph()
            declaration_title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            run = declaration_title.add_run("DECLARATION\n")
            run.bold = True
            run.font.size = Pt(14)
            run.font.color.rgb = RGBColor(255, 0, 0)
            run.underline = True

            declaration = doc.add_paragraph()
            declaration.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

            # 1. Project Title (Blue)
            run = declaration.add_run(
                "We declare that the course-based project work entitled "
            )
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor(0, 0, 0)  # Black

            run = declaration.add_run('"' + title_text + '"')
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor(0, 0, 255)  # Blue
            run.bold = True

            run = declaration.add_run("submitted in the ")
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor(0, 0, 0)  # Black

            # 2. Department Name (Red)
            run = declaration.add_run("Department of " + professor_department + ", ")
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor(255, 0, 0)  # Red

            run = declaration.add_run(
                "Vallurupalli Nageswara Rao Vignana Jyothi Institute of Engineering and Technology, Hyderabad, in partial fulfillment of the requirement for the award of the degree of"            )
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor(0, 0, 0)  # Black

            # 3. Degree Name (Blue)
            run = declaration.add_run("Bachelor of Technology in " + department + ", ")
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor(0, 0, 255)  # Blue
            run.bold = True

            run = declaration.add_run(
                "is a bonafide record of our own work carried out under the supervision of "
            )
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor(0, 0, 0)  # Black

            # Create supervisors list conditionally
            supervisors = [f"{main_professor}, {main_professor_designation}, Department of {professor_department}, VNRVJIET"]
            
            # Only add secondary professor if provided and not empty
            if secondary_professor and secondary_professor.strip():
                supervisors.append(f"{secondary_professor}, {secondary_professor_designation}, Department of {professor_department}, VNRVJIET")

            # Use supervisors list in the document
            for i, supervisor in enumerate(supervisors):
                run = declaration.add_run(str(supervisor))  # Convert to string explicitly
                run.font.size = Pt(12)
                run.font.color.rgb = RGBColor(0, 128, 0)  # Green
                if i < len(supervisors) - 1:
                    run = declaration.add_run(" and ")  # Add "and" between supervisors only if there are multiple

            # ______________________________________________Increase Word Spacing for the entire paragraph_______________________________________________________________________________________________
            declaration.paragraph_format.word_spacing = 2.5  # 125% of normal spacing

            run = declaration.add_run(
                ". Also, we declare that the matter embodied in this thesis has not been submitted by us in full or in any part thereof for the award of any degree of any other institution or university previously.\n\n"
            )
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor(0, 0, 0)  # Black

            # Place
            place = doc.add_paragraph()
            place.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            run = place.add_run("Place: Hyderabad.")
            run.font.size = Pt(12)

            # Create a table
            student_table = doc.add_table(rows=1, cols=4)
            student_table.autofit = False
            make_table_invisible(student_table)

            # Set column widths and add student details
            col_widths = [Inches(1.5), Inches(1.5), Inches(1.5), Inches(1.5)]
            for col, width in zip(student_table.columns, col_widths):
                col.width = width

            students = [
                (member.name for member in team_members),
                (member.rollNumber for member in team_members)
            ]

            for student_row in students:
                row = student_table.add_row()
                for cell, text in zip(row.cells, student_row):
                    cell.text = text
                    for paragraph in cell.paragraphs:
                        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        for run in paragraph.runs:
                            run.font.size = Pt(12)

            # Remove table borders
            make_table_invisible(table)

            # Add some space after the table
            doc.add_paragraph()

            ###############4th page ##########################################################################################################

            doc.add_page_break()
            section = doc.sections[-1]
            add_page_border(section)

            ack_heading = doc.add_paragraph()
            ack_heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            run = ack_heading.add_run("ACKNOWLEDGEMENT")
            run.font.bold = True
            run.font.size = Pt(16)
            run.font.color.rgb = RGBColor(255, 0, 0)  # Red color

            # Add the acknowledgment content
            ack_content = (
                "We express our deep sense of gratitude to our beloved President, "
                "Sri.D.Suresh Babu, VNR Vignana Jyothi Institute of Engineering & Technology for the "
                "valuable guidance and for permitting us to carry out this project.\n\n"
                "With immense pleasure, we record our deep sense of gratitude to our beloved Principal, "
                "Dr.C.D Naidu, for permitting us to carry out this project.\n\n"
                f"We express our deep sense of gratitude to our beloved Professor {main_professor}, "
                f"Professor and Head, Department of {department}, VNR Vignana Jyothi "
                "Institute of Engineering & Technology, Hyderabad-500090 for the valuable guidance and suggestions, "
                "keen interest and through encouragement extended throughout the period of project work.\n\n"
                "We take immense pleasure to express our deep sense of gratitude to our beloved Guide, "
                f"{main_professor}, {main_professor_designation}, Department of {department}, "
                "VNR Vignana Jyothi Institute of Engineering & Technology, Hyderabad, for his/her valuable suggestions "
                "and rare insights, for constant source of encouragement and inspiration throughout my project work.\n\n"
                "We express our thanks to all those who contributed for the successful completion of our project work."
            )
            paragraph = doc.add_paragraph(ack_content)
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            for run in paragraph.runs:
                run.font.size = Pt(12)

            # Add a table for names and roll numbers
            names_and_roll_numbers = [["Name", "Roll Number"]] + [
                [member.name, member.rollNumber] for member in team_members
            ]

            # Create a table with a header row
            table = doc.add_table(rows=1, cols=2)

            # Add data rows
            for name, roll_number in names_and_roll_numbers[0:]:
                row_cells = table.add_row().cells
                name = "               " + name
                roll_number = "                    " + roll_number
                row_cells[0].text = name
                row_cells[1].text = roll_number

            # Apply styling to table rows
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT  # Left-align text
                        for run in paragraph.runs:
                            run.font.size = Pt(12)

            # Remove table borders
            tbl = table._element
            tblBorders = tbl.xpath(".//w:tblBorders")
            for tblBorder in tblBorders:
                tblBorder.getparent().remove(tblBorder)
            
            doc.add_page_break()
            
            # Define styles for titles and body text
            title_style = doc.styles["Title"]
            title_style.font.name = "Calibri"
            title_style.font.size = Pt(20)
            title_style.font.bold = True
            title_style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            body_style = doc.styles["Body Text"]
            body_style.font.name = "Times New Roman"
            body_style.font.size = Pt(14)


            # Define sections with their starting page numbers
            sections = OrderedDict()
            current_page = 5  # TOC will be on page 4
            page_numbers = {}
            
            # Map sections to their page numbers
            for section_name in ["Abstract", "Introduction", "Objectives", "Methodology", "Code"]:
                page_numbers[section_name] = current_page
                current_page += 1
            
            if result and (
                (result.resultImages and len(result.resultImages) > 0) or 
                result.codeOutput or 
                result.aiGeneratedContent
            ):
                page_numbers["Results"] = current_page
                current_page += 1
            
            page_numbers["Conclusion"] = current_page

            # Add sections with their content
            sections["Abstract"] = abstract
            sections["Introduction"] = introduction
            sections["Objectives"] = objectives
            sections["Methodology"] = methodology
            sections["Code"] = project_code
            if "Results" in page_numbers:
                sections["Results"] = result
            sections["Conclusion"] = conclusion

            # Add TOC page (Page 4)
            section = doc.sections[-1]
            add_page_border(section)
            
            # Create TOC title
            toc_title = doc.add_paragraph()
            toc_title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            run = toc_title.add_run("TABLE OF CONTENTS")
            run.font.bold = True
            run.font.size = Pt(16)
            run.font.color.rgb = RGBColor(255, 0, 0)
            
            # Add space after title
            doc.add_paragraph()
            
            # Create TOC table
            toc_table = doc.add_table(rows=len(sections) + 1, cols=3)
            toc_table.style = 'Table Grid'
            toc_table.alignment = WD_TABLE_ALIGNMENT.CENTER
            
            # Set column widths
            toc_table.columns[0].width = Inches(1)    # Chapter number
            toc_table.columns[1].width = Inches(4)    # Title
            toc_table.columns[2].width = Inches(1)    # Page number
            
            # Add headers
            header_cells = toc_table.rows[0].cells
            header_cells[0].text = "Chapter"
            header_cells[1].text = "Title"
            header_cells[2].text = "Page"
            
            # Style header row
            for cell in header_cells:
                paragraph = cell.paragraphs[0]
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                run = paragraph.runs[0]
                run.font.bold = True
                run.font.size = Pt(12)
            
            # Add content rows with correct page numbers
            for idx, (section_name, _) in enumerate(sections.items(), 1):
                cells = toc_table.rows[idx].cells
                
                # Chapter number
                cells[0].text = str(idx)
                cells[0].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                
                # Section name
                cells[1].text = section_name
                cells[1].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                
                # Page number
                cells[2].text = str(page_numbers[section_name])
                cells[2].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                
                # Style the row
                for cell in cells:
                    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.size = Pt(12)
            
            # Add page break after TOC
            doc.add_page_break()

            # Add content sections
            for section_name, section_text in sections.items():
                # Add page border for new section
                section = doc.sections[-1]
                add_page_border(section)

                if section_name != "Results":
                    # Add section title
                    paragraph = doc.add_paragraph()
                    paragraph.style = title_style
                    paragraph.add_run(section_name.upper()).bold = True

                    # Add section content using our formatting function
                    if section_text:
                        format_text_content(doc, section_text)
            
                    # Add page break after each section except conclusion
                    if section_name != "Conclusion":
                        doc.add_page_break()
                else:
                    # Handle Results section
                    _add_results_section(doc, result, title_style, project_code)

            doc.save(report_path)

        # Generate the report using your create_project_report function
        create_project_report(
            abstract,
            introduction,
            objectives,
            methodology,
            data.projectCode,
            data.result if data.result else None,
            conclusion,
            data.department
        )

        # After report is generated, cleanup session images
        if session_id:
            cleanup_session_images(session_id)

        # Return the file
        return FileResponse(
            report_path,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            filename=f"output-report.docx"
        )

    except Exception as e:
        print(f"Error in generate_report: {str(e)}")  # Log the error
        raise HTTPException(
            status_code=500,
            detail=f"Failed to generate report: {str(e)}"
        )

@app.post("/api/generate-ai-content")
async def generate_ai_content(data: ProjectData):
    try:
        # Configure Gemini API
        api_key = os.getenv("GEMINI_API_KEY")
        if not api_key:
            raise HTTPException(status_code=500, detail="GEMINI_API_KEY not set")
        
        genai.configure(api_key=api_key)
        model = genai.GenerativeModel('gemini-pro')
        chat = model.start_chat(history=[])

        # Generate AI analysis
        prompt = f"""Analyze the following code and its output, providing insights about:
        1. The code's functionality and performance
        2. Key patterns or interesting aspects in the output
        3. Potential improvements or optimizations
        4. Any notable technical achievements
        5. within 100 words
        Code: {data.projectCode}
        Output: {data.result.codeOutput if data.result and data.result.codeOutput else 'No output provided'}
        """
        
        ai_analysis = chat.send_message(prompt).text
        
        return {"aiContent": ai_analysis}
    
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))